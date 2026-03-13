"""
周均数据对比分析器 - GUI工具
功能：
1. 文件选择上传
2. 数据处理模式选择（周比/日均）
3. 日期范围选择和均值计算
4. 异常值自定义设定（可按指标设定）
5. AI智能分析生成报告
6. 折线趋势图
7. 完整报告导出
"""

import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from datetime import datetime, timedelta
import json
import os
from openai import OpenAI
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import db_auth


# ============== 页面配置 ==============
st.set_page_config(
    page_title="周均数据对比分析器",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 初始化数据库（应用启动时执行一次）
db_auth.init_db()

# ============== 样式 ==============
_LIGHT_CSS = """
<style>
/* ---- 全局背景 ---- */
[data-testid="stAppViewContainer"] {
    background: linear-gradient(135deg, #e8f0fe 0%, #f0f4ff 40%, #e8f5e9 100%);
    min-height: 100vh;
}
[data-testid="stHeader"] {
    background: rgba(240,244,255,0.6);
    backdrop-filter: blur(12px);
    border-bottom: 1px solid rgba(99,140,255,0.2);
}
/* ---- 侧边栏毛玻璃 ---- */
[data-testid="stSidebar"] {
    background: rgba(255,255,255,0.55) !important;
    backdrop-filter: blur(18px) !important;
    border-right: 1px solid rgba(120,160,255,0.25) !important;
    box-shadow: 4px 0 24px rgba(80,120,220,0.08) !important;
}
[data-testid="stSidebar"] .stMarkdown h2,
[data-testid="stSidebar"] .stMarkdown h3 { color: #2a4bbd; }
/* ---- 主标题 ---- */
.main-header { text-align: center; padding: 1.5rem 0 1rem; }
.main-header h1 {
    font-size: 2.2rem; font-weight: 800;
    background: linear-gradient(90deg, #3a7bd5, #00d2ff, #3a7bd5);
    background-size: 200% auto;
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    background-clip: text; animation: shine 4s linear infinite; margin: 0;
}
@keyframes shine { to { background-position: 200% center; } }
.header-line {
    height: 2px;
    background: linear-gradient(90deg, transparent, #3a7bd5 30%, #00d2ff 70%, transparent);
    margin: 10px auto; width: 60%; border-radius: 2px;
}
.header-subtitle { color: #6b7db3; font-size: 0.85rem; letter-spacing: 2px; margin-top: 4px; }
/* ---- Expander ---- */
[data-testid="stExpander"] {
    background: rgba(255,255,255,0.6) !important;
    backdrop-filter: blur(12px) !important;
    border: 1px solid rgba(200,215,255,0.5) !important;
    border-radius: 12px !important; margin: 8px 0 !important;
    box-shadow: 0 4px 16px rgba(60,100,200,0.06) !important;
}
[data-testid="stExpander"] summary { font-weight: 600 !important; color: #2a4bbd !important; }
/* ---- 按钮 ---- */
.stButton > button {
    background: linear-gradient(135deg, #3a7bd5 0%, #00a8cc 100%) !important;
    color: white !important; border: none !important; border-radius: 10px !important;
    font-weight: 600 !important; box-shadow: 0 4px 14px rgba(58,123,213,0.3) !important;
    transition: all 0.2s !important;
}
.stButton > button:hover {
    box-shadow: 0 6px 20px rgba(58,123,213,0.45) !important; transform: translateY(-1px) !important;
}
/* ---- 输入框 ---- */
[data-testid="stTextArea"] textarea, [data-testid="stTextInput"] input {
    background: rgba(255,255,255,0.75) !important;
    border: 1px solid rgba(120,160,255,0.35) !important; border-radius: 10px !important;
}
[data-testid="stTextArea"] textarea:focus, [data-testid="stTextInput"] input:focus {
    border-color: #3a7bd5 !important; box-shadow: 0 0 0 3px rgba(58,123,213,0.12) !important;
}
/* ---- Metric ---- */
[data-testid="stMetric"] {
    background: rgba(255,255,255,0.65) !important;
    border: 1px solid rgba(200,215,255,0.5) !important; border-radius: 12px !important;
    padding: 12px 16px !important; box-shadow: 0 4px 16px rgba(60,100,200,0.06) !important;
}
[data-testid="stMetricValue"] { color: #2a4bbd !important; }
/* ---- 数据表格 ---- */
[data-testid="stDataFrame"] {
    border-radius: 12px !important; border: 1px solid rgba(200,215,255,0.4) !important;
    box-shadow: 0 4px 16px rgba(60,100,200,0.07) !important; overflow: hidden !important;
}
/* ---- Divider ---- */
hr {
    border: none !important; height: 1px !important;
    background: linear-gradient(90deg, transparent, rgba(58,123,213,0.3), transparent) !important;
    margin: 1rem 0 !important;
}
/* ---- 分区小标题 ---- */
.section-header {
    display: flex; align-items: center; gap: 8px; font-size: 1rem; font-weight: 700;
    color: #2a4bbd; padding: 7px 14px; background: rgba(58,123,213,0.07);
    border-left: 3px solid #3a7bd5; border-radius: 0 8px 8px 0; margin: 14px 0 8px 0;
}
/* ---- 状态标签 ---- */
.abnormal-tag {
    display: inline-block; background: linear-gradient(135deg, #f45b69, #ff8b5e);
    color: white; padding: 2px 10px; border-radius: 20px; font-size: 0.75rem; font-weight: 600;
}
.normal-tag {
    display: inline-block; background: linear-gradient(135deg, #00c6a2, #00e0b6);
    color: white; padding: 2px 10px; border-radius: 20px; font-size: 0.75rem; font-weight: 600;
}
/* ---- 统计卡片 ---- */
.metric-card {
    background: rgba(255,255,255,0.7); backdrop-filter: blur(10px);
    border: 1px solid rgba(255,255,255,0.9); border-radius: 14px; padding: 16px 20px;
    text-align: center; box-shadow: 0 4px 20px rgba(60,100,200,0.08);
    position: relative; overflow: hidden;
}
.metric-card::before {
    content: ''; position: absolute; top: 0; left: 0; right: 0; height: 3px;
    background: linear-gradient(90deg, #3a7bd5, #00d2ff); border-radius: 14px 14px 0 0;
}
</style>
"""

_DARK_CSS = """
<style>
/* ---- 背景 ---- */
[data-testid="stAppViewContainer"] {
    background: linear-gradient(135deg, #0d1117 0%, #0f1724 50%, #0d1a14 100%) !important;
}
[data-testid="stHeader"] {
    background: rgba(13,17,23,0.85) !important;
    border-bottom: 1px solid rgba(99,140,255,0.15) !important;
}
[data-testid="stSidebar"] {
    background: rgba(22,27,36,0.95) !important;
    border-right: 1px solid rgba(99,140,255,0.15) !important;
    box-shadow: 4px 0 24px rgba(0,0,0,0.4) !important;
}
/* ---- 全局文字（base=light时覆盖）---- */
.block-container, .block-container p,
.block-container h1, .block-container h2,
.block-container h3, .block-container h4,
.block-container h5, .block-container h6 {
    color: #c9d1d9 !important;
}
h1, h2, h3, h4, h5, h6 { color: #e6edf3 !important; }
[data-testid="stMarkdownContainer"],
[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li,
[data-testid="stMarkdownContainer"] span { color: #c9d1d9 !important; }
/* label：选择框/单选/复选/滑块 */
label, [data-testid="stWidgetLabel"],
[data-testid="stSlider"] label,
[data-testid="stRadio"] label,
[data-testid="stCheckbox"] label,
[data-testid="stSelectbox"] label,
[data-testid="stMultiSelect"] label,
[data-testid="stFileUploader"] label { color: #c9d1d9 !important; }
/* 侧边栏文字 */
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] .stMarkdown h2,
[data-testid="stSidebar"] .stMarkdown h3 { color: #79a8ff !important; }
/* ---- Expander ---- */
[data-testid="stExpander"] {
    background: rgba(22,27,36,0.85) !important;
    border: 1px solid rgba(99,140,255,0.18) !important;
    box-shadow: 0 4px 16px rgba(0,0,0,0.3) !important;
}
[data-testid="stExpander"] summary,
[data-testid="stExpander"] summary p { color: #79a8ff !important; }
/* ---- 输入框 ---- */
[data-testid="stTextArea"] textarea, [data-testid="stTextInput"] input {
    background: rgba(30,37,48,0.9) !important;
    border: 1px solid rgba(99,140,255,0.25) !important;
    color: #c9d1d9 !important;
}
/* ---- Metric ---- */
[data-testid="stMetric"] {
    background: rgba(22,27,36,0.85) !important;
    border: 1px solid rgba(99,140,255,0.2) !important;
    box-shadow: 0 4px 16px rgba(0,0,0,0.3) !important;
}
[data-testid="stMetricValue"] { color: #79a8ff !important; }
[data-testid="stMetricLabel"] { color: #8b949e !important; }
/* ---- 数据表格 ---- */
[data-testid="stDataFrame"] {
    border: 1px solid rgba(99,140,255,0.2) !important;
    box-shadow: 0 4px 16px rgba(0,0,0,0.3) !important;
}
/* ---- 分区标题/卡片 ---- */
.section-header {
    color: #79a8ff !important;
    background: rgba(58,123,213,0.12) !important;
    border-left-color: #4a90d9 !important;
}
.metric-card {
    background: rgba(22,27,36,0.85) !important;
    border: 1px solid rgba(99,140,255,0.2) !important;
    box-shadow: 0 4px 20px rgba(0,0,0,0.35) !important;
}
.header-subtitle { color: #8b949e !important; }
/* ---- caption / help text ---- */
[data-testid="stCaptionContainer"] p,
small, .stCaption { color: #8b949e !important; }
</style>
"""

st.markdown(_LIGHT_CSS, unsafe_allow_html=True)
if st.session_state.get("dark_mode", False):
    st.markdown(_DARK_CSS, unsafe_allow_html=True)


# ============== 数据处理函数 ==============
def detect_format(series):
    """检测列的原始格式（百分比、小数位数）"""
    format_info = {'is_percent': False, 'decimal_places': 0}
    
    for val in series.dropna().head(10):
        val_str = str(val).strip()
        if '%' in val_str:
            format_info['is_percent'] = True
            # 检测百分比的小数位数
            num_part = val_str.replace('%', '').replace(',', '')
            if '.' in num_part:
                format_info['decimal_places'] = max(format_info['decimal_places'], len(num_part.split('.')[-1]))
        elif '.' in val_str:
            try:
                num_part = val_str.replace(',', '')
                format_info['decimal_places'] = max(format_info['decimal_places'], len(num_part.split('.')[-1]))
            except:
                pass
    
    return format_info


def format_value(val, format_info):
    """根据格式信息格式化数值"""
    if pd.isna(val):
        return val
    
    decimal_places = format_info.get('decimal_places', 2)
    is_percent = format_info.get('is_percent', False)
    
    if is_percent:
        # 百分比格式：内部存储为小数，显示时转换为百分比
        return round(val * 100, decimal_places)
    else:
        return round(val, decimal_places)


def clean_numeric_value(val):
    """清洗数值，保留原始值用于计算"""
    if pd.isna(val) or val == '' or val == 'nan':
        return np.nan
    val_str = str(val).strip().replace(',', '').replace('，', '')
    if '%' in val_str:
        try:
            return float(val_str.replace('%', '')) / 100
        except:
            return val
    try:
        return float(val_str)
    except:
        return val


def clean_data(df):
    """清洗数据"""
    cleaned = df.copy()
    cleaned.columns = cleaned.columns.str.strip()
    
    for col in cleaned.columns:
        if cleaned[col].dtype == 'object':
            cleaned[col] = cleaned[col].apply(clean_numeric_value)
    
    return cleaned


def calculate_comparison(df, metric_col, prev_col, curr_col, thresholds, prev_label="上周均值", curr_label="本周均值", format_info=None):
    """计算对比数据，保持原始数据的小数位数格式"""
    result = pd.DataFrame()
    result['指标'] = df[metric_col]
    
    # 移除_format列（如果存在）用于显示
    if '_format' in df.columns:
        df_clean = df.drop(columns=['_format'])
    else:
        df_clean = df
    
    prev_values = pd.to_numeric(df_clean[prev_col], errors='coerce').fillna(0)
    curr_values = pd.to_numeric(df_clean[curr_col], errors='coerce').fillna(0)
    diff_values = curr_values - prev_values
    
    # 根据每个指标的格式信息格式化数值
    def format_with_decimals(value, metric_name, add_sign=False):
        """根据格式信息格式化数值，保持原始格式（百分比、小数位数）"""
        decimals = 0  # 默认无小数
        is_percent = False
        
        if format_info and metric_name in format_info:
            decimals = format_info[metric_name].get('decimal_places', 0)
            is_percent = format_info[metric_name].get('is_percent', False)
        
        # 百分比数据需要显示%号
        if is_percent:
            if add_sign and value > 0:
                return f"+{value:.{decimals}f}%"
            else:
                return f"{value:.{decimals}f}%"
        else:
            if add_sign and value > 0:
                return f"+{value:.{decimals}f}"
            else:
                return f"{value:.{decimals}f}"
    
    # 格式化上周、本周、差值列
    formatted_prev = []
    formatted_curr = []
    formatted_diff = []
    
    for i, metric in enumerate(df[metric_col]):
        formatted_prev.append(format_with_decimals(prev_values.iloc[i], metric))
        formatted_curr.append(format_with_decimals(curr_values.iloc[i], metric))
        formatted_diff.append(format_with_decimals(diff_values.iloc[i], metric, add_sign=True))
    
    result[prev_label] = formatted_prev
    result[curr_label] = formatted_curr
    result['差值'] = formatted_diff
    result['_diff_raw'] = diff_values  # 保留原始数值用于计算
    result['_prev_raw'] = prev_values  # 保留原始数值用于计算
    
    def calc_rate(row):
        try:
            prev = row['_prev_raw']
            diff = row['_diff_raw']
            if pd.notna(prev) and prev != 0:
                rate = (diff / prev) * 100
                if pd.isna(rate) or not np.isfinite(rate):
                    return 0.0
                return round(float(rate), 2)
        except Exception:
            pass
        return 0.0
    
    result['_rate_raw'] = result.apply(calc_rate, axis=1)  # 保留原始数值
    # 格式化涨跌率显示正负号（始终保证百分比格式）
    def _fmt_rate(x):
        try:
            v = float(x)
            if not np.isfinite(v):
                return "0.00%"
            return f"+{v:.2f}%" if v > 0 else f"{v:.2f}%"
        except Exception:
            return "0.00%"
    result['涨跌率(%)'] = result['_rate_raw'].apply(_fmt_rate)
    
    # 根据自定义阈值判断异常
    def check_abnormal(row):
        metric = row['指标']
        rate = abs(row['_rate_raw'])
        threshold = thresholds.get(metric, thresholds.get('__default__', 15))
        return rate >= threshold
    
    result['是否异常'] = result.apply(check_abnormal, axis=1)
    result['状态'] = result.apply(
        lambda row: f"{'↑' if row['_rate_raw'] > 0 else '↓'} {'异常' if row['是否异常'] else '正常'}" 
        if abs(row['_rate_raw']) >= 1 else '稳定', 
        axis=1
    )
    
    # 移除内部计算列
    result = result.drop(columns=['_diff_raw', '_rate_raw', '_prev_raw'])
    
    return result


def parse_date_with_weekday(date_str):
    """解析可能包含星期标注的日期，如 '2026-01-11(日)' 或 '2026-01-10(六)'"""
    if pd.isna(date_str):
        return pd.NaT
    
    date_str = str(date_str).strip()
    
    # 移除括号中的星期标注，如 (六)、(日)、(一) 等
    import re
    cleaned = re.sub(r'\([一二三四五六日]\)', '', date_str)
    cleaned = re.sub(r'\(周[一二三四五六日]\)', '', cleaned)
    cleaned = re.sub(r'\(星期[一二三四五六日]\)', '', cleaned)
    cleaned = cleaned.strip()
    
    try:
        return pd.to_datetime(cleaned)
    except:
        return pd.NaT


def calculate_daily_average(df, date_col, metric_cols, start_date, end_date, original_df=None):
    """计算日期范围内的日均值（包含所有日期，包括周末）"""
    df = df.copy()
    # 使用自定义解析函数处理带星期标注的日期
    df[date_col] = df[date_col].apply(parse_date_with_weekday)
    
    # 筛选日期范围（包含起始和结束日期，包括周末）
    start_dt = pd.to_datetime(start_date)
    end_dt = pd.to_datetime(end_date) + pd.Timedelta(days=0, hours=23, minutes=59, seconds=59)
    mask = (df[date_col] >= start_dt) & (df[date_col] <= end_dt)
    filtered_df = df[mask]
    
    # 计算日期范围的实际天数（包含起止日期）
    date_range_days = (pd.to_datetime(end_date) - pd.to_datetime(start_date)).days + 1
    
    if filtered_df.empty:
        return None, 0, {}, date_range_days
    
    # 检测原始数据格式
    format_info = {}
    if original_df is not None:
        for col in metric_cols:
            if col in original_df.columns:
                format_info[col] = detect_format(original_df[col])
            else:
                format_info[col] = {'is_percent': False, 'decimal_places': 0}
    
    # 计算均值
    result = {}
    for col in metric_cols:
        values = pd.to_numeric(filtered_df[col], errors='coerce')
        avg_val = values.mean()
        # 根据原始格式处理
        if col in format_info:
            result[col] = format_value(avg_val, format_info[col])
        else:
            result[col] = avg_val
    
    data_records = len(filtered_df)
    return result, data_records, format_info, date_range_days


def create_trend_chart(df, date_col, metric_cols, title="数据趋势图"):
    """创建折线趋势图"""
    df = df.copy()
    # 使用自定义解析函数处理带星期标注的日期
    df[date_col] = df[date_col].apply(parse_date_with_weekday)
    df = df.sort_values(date_col)
    
    fig = go.Figure()
    
    for col in metric_cols:
        values = pd.to_numeric(df[col], errors='coerce')
        fig.add_trace(go.Scatter(
            x=df[date_col],
            y=values,
            mode='lines+markers',
            name=col,
            hovertemplate=f'{col}: %{{y:,.2f}}<extra></extra>'
        ))
    
    fig.update_layout(
        title=title,
        xaxis_title="日期",
        yaxis_title="数值",
        hovermode='x unified',
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
    )
    
    return fig


def create_comparison_bar_chart(comparison_df, metric_col, prev_label, curr_label):
    """创建对比柱状图"""
    def to_numeric_col(col):
        return pd.to_numeric(
            col.astype(str).str.replace('%', '').str.replace('+', '').str.replace(',', ''),
            errors='coerce'
        ).fillna(0)

    fig = go.Figure()
    
    fig.add_trace(go.Bar(
        name=prev_label,
        x=comparison_df[metric_col],
        y=to_numeric_col(comparison_df[prev_label]),
        marker_color='rgba(58,123,213,0.75)',
        marker_line_color='rgba(58,123,213,1)',
        marker_line_width=1,
    ))
    
    fig.add_trace(go.Bar(
        name=curr_label,
        x=comparison_df[metric_col],
        y=to_numeric_col(comparison_df[curr_label]),
        marker_color='rgba(0,210,255,0.75)',
        marker_line_color='rgba(0,180,220,1)',
        marker_line_width=1,
    ))
    
    fig.update_layout(
        title="周期对比柱状图",
        barmode='group',
        xaxis_title="指标",
        yaxis_title="数值",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(255,255,255,0.4)",
        font=dict(color="#3a4d7a"),
        legend=dict(
            orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
            bgcolor="rgba(255,255,255,0.6)", bordercolor="rgba(120,160,255,0.3)", borderwidth=1
        ),
        xaxis=dict(gridcolor="rgba(120,160,255,0.1)"),
        yaxis=dict(gridcolor="rgba(120,160,255,0.1)"),
    )
    
    return fig


def create_change_rate_chart(comparison_df, thresholds):
    """创建涨跌率图表"""
    df = comparison_df.copy()
    
    # 解析涨跌率字符串为数值（去掉%和+号）
    def parse_rate(rate_str):
        try:
            return float(str(rate_str).replace('%', '').replace('+', ''))
        except:
            return 0
    
    rate_values = df['涨跌率(%)'].apply(parse_rate)
    
    # 根据异常状态设置颜色
    colors = ['rgba(244,91,105,0.8)' if abnormal else 'rgba(0,184,124,0.8)' for abnormal in df['是否异常']]

    fig = go.Figure()

    fig.add_trace(go.Bar(
        x=df['指标'],
        y=rate_values,
        marker_color=colors,
        marker_line_width=0,
        text=df['涨跌率(%)'],
        textposition='outside'
    ))

    # 添加阈值线
    default_threshold = thresholds.get('__default__', 15)
    fig.add_hline(y=default_threshold, line_dash="dash", line_color="rgba(244,91,105,0.6)",
                  annotation_text=f"异常阈值 +{default_threshold}%")
    fig.add_hline(y=-default_threshold, line_dash="dash", line_color="rgba(244,91,105,0.6)",
                  annotation_text=f"异常阈值 -{default_threshold}%")
    fig.add_hline(y=0, line_color="rgba(58,123,213,0.3)", line_width=1)

    fig.update_layout(
        title="涨跌率分布（红色=异常，绿色=正常）",
        xaxis_title="指标",
        yaxis_title="涨跌率(%)",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(255,255,255,0.4)",
        font=dict(color="#3a4d7a"),
        xaxis=dict(gridcolor="rgba(120,160,255,0.1)"),
        yaxis=dict(gridcolor="rgba(120,160,255,0.1)", ticksuffix="%"),
    )
    
    return fig


def generate_ai_analysis(comparison_df, original_df, thresholds, api_key, base_url=None):
    """使用AI生成分析报告"""
    if not api_key:
        return generate_rule_based_analysis(comparison_df, thresholds)
    
    try:
        client_kwargs = {"api_key": api_key}
        if base_url:
            client_kwargs["base_url"] = base_url
        
        client = OpenAI(**client_kwargs)
        
        # 准备数据摘要
        data_summary = comparison_df.to_string(index=False)
        abnormal_data = comparison_df[comparison_df['是否异常'] == True]
        abnormal_summary = abnormal_data.to_string(index=False) if not abnormal_data.empty else "无异常数据"
        
        threshold_info = "\n".join([f"- {k}: {v}%" for k, v in thresholds.items()])
        
        prompt = f"""你是一位资深的数据分析师，请对以下周均数据对比进行深度分析。

## 数据概览
{data_summary}

## 异常数据（超过设定阈值）
{abnormal_summary}

## 异常阈值设定
{threshold_info}

请提供以下分析：

### 1. 整体趋势分析
分析本周相比上周的整体变化趋势

### 2. 异常数据分析
针对每个异常指标：
- 异常程度评估
- 可能的原因分析
- 是否属于正常波动范围

### 3. 漏斗分析（如适用）
如果数据呈现漏斗结构（如：曝光→点击→转化），分析各环节转化情况

### 4. 关注建议
- 需要重点关注的环节
- 建议采取的行动

### 5. 总结结论
用2-3句话总结本周数据表现

请用专业但易懂的语言，给出具体可操作的建议。"""

        # 根据base_url选择模型
        if "openrouter" in (base_url or "").lower():
            # 从session_state获取用户选择的模型
            model = st.session_state.get('selected_model', 'openai/gpt-4o-mini')
        else:
            model = "gpt-4o-mini"
        
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你是一位专业的数据分析师，擅长解读业务数据并给出洞察。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.warning(f"AI分析失败: {str(e)}，使用规则分析")
        return generate_rule_based_analysis(comparison_df, thresholds)


def generate_rule_based_analysis(comparison_df, thresholds):
    """基于规则的分析（无AI时使用）"""
    report = []
    
    # 1. 整体趋势
    report.append("### 1. 整体趋势分析\n")
    avg_change = comparison_df['涨跌率(%)'].mean()
    if avg_change > 5:
        report.append(f"本周整体呈**上升趋势**，平均涨幅 {avg_change:.1f}%。\n")
    elif avg_change < -5:
        report.append(f"本周整体呈**下降趋势**，平均降幅 {abs(avg_change):.1f}%。\n")
    else:
        report.append(f"本周整体表现**平稳**，平均变化 {avg_change:.1f}%。\n")
    
    # 2. 异常分析
    report.append("\n### 2. 异常数据分析\n")
    abnormal = comparison_df[comparison_df['是否异常'] == True]
    if abnormal.empty:
        report.append("本周无异常数据，各指标变化均在正常范围内。\n")
    else:
        for _, row in abnormal.iterrows():
            metric = row['指标']
            rate = row['涨跌率(%)']
            threshold = thresholds.get(metric, thresholds.get('__default__', 15))
            direction = "上涨" if rate > 0 else "下降"
            
            report.append(f"**{metric}**：{direction} {abs(rate):.1f}%（阈值 {threshold}%）\n")
            report.append(f"  - 异常程度：{'严重' if abs(rate) > threshold * 2 else '中等'}\n")
            report.append(f"  - 可能原因：需结合业务背景进一步分析\n\n")
    
    # 3. 关注建议
    report.append("\n### 3. 关注建议\n")
    if not abnormal.empty:
        report.append("建议重点关注以下指标：\n")
        for metric in abnormal['指标'].tolist():
            report.append(f"  - {metric}\n")
    else:
        report.append("各指标表现正常，建议保持现有运营策略。\n")
    
    # 4. 总结
    report.append("\n### 4. 总结结论\n")
    abnormal_count = len(abnormal)
    total_count = len(comparison_df)
    report.append(f"本周共监测 {total_count} 项指标，其中 {abnormal_count} 项出现异常波动。")
    
    if abnormal_count == 0:
        report.append("整体运营状况良好，数据表现稳定。")
    elif abnormal_count <= 2:
        report.append("存在局部波动，建议针对性排查原因。")
    else:
        report.append("多项指标异常，建议全面复盘本周运营情况。")
    
    return "\n".join(report)


def create_excel_report(comparison_df, analysis_text, thresholds):
    """创建Excel报告（只包含处理后的周均数据）"""
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Sheet 1: 周均对比分析
        comparison_df.to_excel(writer, sheet_name='1_周均对比分析', index=False)
        
        # Sheet 2: 异常数据
        abnormal_df = comparison_df[comparison_df['是否异常'] == True]
        if not abnormal_df.empty:
            abnormal_df.to_excel(writer, sheet_name='2_异常数据', index=False)
        
        # Sheet 3: 阈值设定
        threshold_df = pd.DataFrame([
            {'指标': k, '异常阈值(%)': v} 
            for k, v in thresholds.items() if k != '__default__'
        ])
        # 添加默认阈值
        default_row = pd.DataFrame([{'指标': '默认阈值', '异常阈值(%)': thresholds.get('__default__', 15)}])
        threshold_df = pd.concat([default_row, threshold_df], ignore_index=True)
        threshold_df.to_excel(writer, sheet_name='3_阈值设定', index=False)
        
        # Sheet 4: 分析结论
        analysis_df = pd.DataFrame({'分析报告': [analysis_text]})
        analysis_df.to_excel(writer, sheet_name='4_分析结论', index=False)
    
    output.seek(0)
    return output


def create_markdown_report(comparison_df, analysis_text, thresholds):
    """创建Markdown报告（只包含处理后的周均数据）"""
    report = []
    report.append("# 周均数据对比分析报告\n")
    report.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    report.append("---\n")
    
    # 1. 周均对比分析
    report.append("\n## 一、周均对比分析\n")
    report.append(comparison_df.to_markdown(index=False))
    
    # 2. 异常数据
    abnormal_df = comparison_df[comparison_df['是否异常'] == True]
    if not abnormal_df.empty:
        report.append("\n\n## 二、异常数据\n")
        report.append(abnormal_df.to_markdown(index=False))
    
    # 3. 异常阈值设定
    report.append("\n\n## 三、异常阈值设定\n")
    report.append(f"- **默认阈值**: {thresholds.get('__default__', 15)}%\n")
    for k, v in thresholds.items():
        if k != "__default__":
            report.append(f"- **{k}**: {v}%\n")
    
    # 4. 分析结论
    report.append("\n\n## 四、分析结论\n")
    report.append(analysis_text)
    
    report.append("\n\n---\n")
    report.append("*报告由周均数据对比分析器自动生成*")
    
    return "\n".join(report)


def create_word_report(comparison_df, analysis_text, thresholds):
    """创建Word报告（只包含处理后的周均数据）"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('周均数据对比分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 生成时间
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # 一、周均对比分析
    doc.add_heading('一、周均对比分析', level=1)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=len(comparison_df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    header_cells = table.rows[0].cells
    for i, col in enumerate(comparison_df.columns):
        header_cells[i].text = str(col)
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # 数据行
    for _, row in comparison_df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    
    doc.add_paragraph()
    
    # 二、异常数据
    abnormal_df = comparison_df[comparison_df['是否异常'] == True]
    if not abnormal_df.empty:
        doc.add_heading('二、异常数据', level=1)
        
        abnormal_table = doc.add_table(rows=1, cols=len(abnormal_df.columns))
        abnormal_table.style = 'Table Grid'
        abnormal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_cells = abnormal_table.rows[0].cells
        for i, col in enumerate(abnormal_df.columns):
            header_cells[i].text = str(col)
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        for _, row in abnormal_df.iterrows():
            row_cells = abnormal_table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        
        doc.add_paragraph()
    
    # 三、异常阈值设定
    doc.add_heading('三、异常阈值设定', level=1)
    doc.add_paragraph(f"• 默认阈值: {thresholds.get('__default__', 15)}%")
    for k, v in thresholds.items():
        if k != "__default__":
            doc.add_paragraph(f"• {k}: {v}%")
    
    doc.add_paragraph()
    
    # 四、分析结论
    doc.add_heading('四、分析结论', level=1)
    
    # 解析分析文本，按段落添加
    for line in analysis_text.split('\n'):
        line = line.strip()
        if line:
            if line.startswith('###'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.replace('**', ''))
                run.bold = True
            else:
                doc.add_paragraph(line)
    
    doc.add_paragraph()
    doc.add_paragraph("---")
    footer = doc.add_paragraph("报告由周均数据对比分析器自动生成")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存到BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ============== 登录 / 注册页面 ==============
def show_auth_page():
    st.markdown("""
    <div style="max-width:420px;margin:80px auto 0;padding:36px 40px 32px;
                background:rgba(255,255,255,0.72);backdrop-filter:blur(16px);
                border-radius:20px;border:1px solid rgba(120,160,255,0.25);
                box-shadow:0 8px 32px rgba(58,123,213,0.12);">
        <div style="text-align:center;margin-bottom:28px;">
            <svg width="48" height="48" viewBox="0 0 48 48" fill="none" xmlns="http://www.w3.org/2000/svg">
                <rect width="48" height="48" rx="14" fill="url(#authg)"/>
                <circle cx="24" cy="19" r="7" fill="white" fill-opacity="0.9"/>
                <path d="M10 40c0-7.732 6.268-14 14-14s14 6.268 14 14" stroke="white"
                      stroke-width="2.5" stroke-linecap="round" fill="none" fill-opacity="0.9"/>
                <defs><linearGradient id="authg" x1="0" y1="0" x2="48" y2="48" gradientUnits="userSpaceOnUse">
                    <stop offset="0%" stop-color="#3a7bd5"/>
                    <stop offset="100%" stop-color="#00d2ff"/>
                </linearGradient></defs>
            </svg>
            <h2 style="margin:12px 0 4px;color:#1e3a6e;font-size:1.4rem;font-weight:700;">周均数据对比分析器</h2>
            <p style="color:#6b7db3;font-size:0.85rem;margin:0;">请登录或注册以使用全部功能</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    _, col, _ = st.columns([1, 1.6, 1])
    with col:
        tab_login, tab_reg = st.tabs(["🔑 登录", "📝 注册"])

        with tab_login:
            with st.form("login_form"):
                u = st.text_input("用户名", key="login_u")
                p = st.text_input("密码", type="password", key="login_p")
                if st.form_submit_button("登录", use_container_width=True, type="primary"):
                    ok, msg, uid = db_auth.login_user(u, p)
                    if ok:
                        st.session_state["authenticated"] = True
                        st.session_state["username"] = u.strip()
                        st.session_state["user_id"] = uid
                        st.rerun()
                    else:
                        st.error(msg)

        with tab_reg:
            with st.form("reg_form"):
                u2 = st.text_input("用户名", key="reg_u")
                p2 = st.text_input("密码（至少 6 位）", type="password", key="reg_p")
                p3 = st.text_input("确认密码", type="password", key="reg_p2")
                if st.form_submit_button("注册", use_container_width=True, type="primary"):
                    if p2 != p3:
                        st.error("两次密码不一致")
                    else:
                        ok, msg = db_auth.register_user(u2, p2)
                        if ok:
                            st.success(msg)
                        else:
                            st.error(msg)


# ============== 示例数据 ==============
def get_demo_comparison():
    """返回示例对比数据，列结构与真实 comparison_df 完全一致"""
    data = {
        '指标':    ['DAU', 'MAU', '次日留存', '7日留存', '付费率', 'ARPU', 'LTV', '人均时长',
                   'DAU/MAU', '新增用户', '付费金额', '活跃天数'],
        '上周均值': ['125000', '860000', '32.1', '18.5', '5.20%', '18.50', '220.0', '42.3',
                   '14.5', '8500', '42000', '3.8'],
        '本周均值': ['133000', '895000', '34.5', '19.8', '4.40%', '22.00', '235.0', '44.1',
                   '11.6', '9200', '44500', '3.9'],
        '差值':    ['+8000', '+35000', '+2.4', '+1.3', '-0.80%', '+3.50', '+15.0', '+1.8',
                   '-2.9', '+700', '+2500', '+0.1'],
        '涨跌率(%)': ['+6.40%', '+4.07%', '+7.48%', '+7.03%', '-15.38%', '+18.92%', '+6.82%', '+4.26%',
                    '-20.00%', '+8.24%', '+5.95%', '+2.63%'],
        '是否异常': [False, False, False, False, True, True, False, False,
                   True, False, False, True],
        '状态':    ['↑ 正常', '↑ 正常', '↑ 正常', '↑ 正常', '↓ 异常', '↑ 异常', '↑ 正常', '↑ 正常',
                   '↓ 异常', '↑ 正常', '↑ 正常', '↑ 异常'],
    }
    return pd.DataFrame(data)


# ============== 月度核心数据看板 ==============
def show_monthly_dashboard(uid):
    import plotly.graph_objects as _go_m

    st.markdown("#### 📅 月度核心数据看板")
    st.caption("上传日粒度数据，按月聚合求和与均值，支持月份范围筛选")

    # ── 文件上传 ──
    mf = st.file_uploader(
        "上传数据文件（CSV / Excel，每行一天）",
        type=["csv", "xlsx", "xls"],
        key="monthly_uploader",
        help="日期列 + 各指标列，每行为一天数据"
    )

    if mf is None:
        st.info("请上传数据文件，格式示例：\n\n| 日期 | 指标A | 指标B |\n|------|-------|-------|\n| 2026-01-01 | 5000 | 120 |\n| 2026-01-02 | 4800 | 115 |")
        return

    # ── 读取 ──
    try:
        if mf.name.endswith(".csv"):
            raw = None
            for enc in ["utf-8-sig", "utf-8", "gbk"]:
                try:
                    mf.seek(0); raw = pd.read_csv(mf, encoding=enc); break
                except Exception:
                    continue
            if raw is None:
                st.error("CSV 编码识别失败"); return
            mdf = raw
        else:
            mf.seek(0); mdf = pd.read_excel(mf)
    except Exception as e:
        st.error(f"读取失败：{e}"); return

    mcols = list(mdf.columns)
    st.success(f"✅ 已加载 {len(mdf)} 行 × {len(mcols)} 列")

    # ── 列配置 ──
    st.markdown("##### 🔧 列配置")
    date_col = st.selectbox("日期列", mcols, index=0, key="m_date_col")
    non_date = [c for c in mcols if c != date_col]

    st.markdown("**指标列**")
    if st.session_state.pop("_m_uncheck_selectall", False):
        st.session_state["m_select_all"] = False
    m_select_all = st.checkbox("全选所有指标", value=True, key="m_select_all")
    if m_select_all:
        metric_cols = non_date
    else:
        _m_tpl_default = st.session_state.pop("_m_tpl_load", non_date)
        metric_cols = st.multiselect(
            "选择指标列", non_date,
            default=[m for m in _m_tpl_default if m in non_date] or non_date,
            key="m_metric_cols"
        )

    # 指标模板管理（月度看板）
    with st.expander("📋 指标模板管理", expanded=False):
        _mmt_list = db_auth.get_metric_templates(uid)
        _mmt_names = [t["template_name"] for t in _mmt_list]
        if _mmt_list:
            st.markdown("**已保存的模板**")
            _mmt_sel = st.selectbox("选择模板", _mmt_names, key="mmt_select")
            _mmt_data = next((t for t in _mmt_list if t["template_name"] == _mmt_sel), None)
            if _mmt_data:
                _mmt_prev = _mmt_data["metrics"]
                st.caption(f"包含 {len(_mmt_prev)} 个指标：{', '.join(_mmt_prev[:6])}{'…' if len(_mmt_prev)>6 else ''}")
                st.caption(f"📅 更新于：{_mmt_data['updated_at']}")
            _mmt_c1, _mmt_c2, _mmt_c3 = st.columns(3)
            with _mmt_c1:
                if st.button("📥 加载", use_container_width=True, key="mmt_load"):
                    st.session_state["_m_uncheck_selectall"] = True
                    st.session_state["_m_tpl_load"] = _mmt_data["metrics"]
                    st.success(f"已加载「{_mmt_sel}」")
                    st.rerun()
            with _mmt_c2:
                _mmt_rn = st.text_input("重命名为", key="mmt_rename_input", placeholder="输入新名称")
                if st.button("✏️ 重命名", use_container_width=True, key="mmt_rename"):
                    ok, msg = db_auth.rename_metric_template(uid, _mmt_sel, _mmt_rn)
                    st.success(msg) if ok else st.error(msg)
                    st.rerun()
            with _mmt_c3:
                if st.button("🗑️ 删除", use_container_width=True, key="mmt_del"):
                    ok, msg = db_auth.delete_metric_template(uid, _mmt_sel)
                    st.success(msg) if ok else st.error(msg)
                    st.rerun()
        else:
            st.caption("暂无已保存的指标模板")
        st.divider()
        st.markdown("**保存当前选中指标为模板**")
        _mmt_save_name = st.text_input("模板名称", key="mmt_save_name", placeholder="例：月度核心指标")
        if st.button("💾 保存当前指标", use_container_width=True, type="primary", key="mmt_save"):
            ok, msg = db_auth.save_metric_template(uid, _mmt_save_name, list(metric_cols))
            st.success(msg) if ok else st.error(msg)
            if ok:
                st.rerun()

    if not metric_cols:
        st.warning("请至少选择一个指标列"); return

    # ── 日期解析 + 数值清洗 ──
    # 先把原始日期字符串存起来（含（六）（日）等标注）
    mdf["_orig_date"] = mdf[date_col].astype(str).str.strip()
    # 仅提取 YYYY-MM-DD 部分用于计算，兼容"2026-02-08（日）"等格式
    import re as _re_date
    mdf[date_col] = pd.to_datetime(
        mdf[date_col].astype(str).str.extract(r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})')[0],
        errors="coerce"
    )
    mdf = mdf.dropna(subset=[date_col])
    mdf = mdf.sort_values(date_col).reset_index(drop=True)
    # 排序后用原始字符串构建预览 DataFrame（保留百分号、星期标注等）
    orig_mdf = mdf[["_orig_date"] + metric_cols].copy()
    orig_mdf = orig_mdf.rename(columns={"_orig_date": date_col})
    mdf = mdf.drop(columns=["_orig_date"])
    # 检测含百分号的列
    pct_cols = set()
    for c in metric_cols:
        sample = orig_mdf[c].dropna().astype(str).head(20)
        if any("%" in v for v in sample):
            pct_cols.add(c)
    for c in metric_cols:
        mdf[c] = pd.to_numeric(
            mdf[c].astype(str).str.replace(",", "").str.replace("%", ""),
            errors="coerce"
        )
    if len(mdf) == 0:
        st.error("日期解析后数据为空，请检查日期列格式"); return

    # ── 月份范围选择 ──
    mdf["_ym"] = mdf[date_col].dt.to_period("M")
    all_months = sorted(mdf["_ym"].unique().tolist())
    month_strs = [str(m) for m in all_months]

    st.markdown("##### 📅 月份范围")
    rc1, rc2 = st.columns(2)
    with rc1:
        start_m = st.selectbox("起始月份", month_strs, index=0, key="m_start")
    with rc2:
        end_m   = st.selectbox("结束月份", month_strs, index=len(month_strs)-1, key="m_end")

    if start_m > end_m:
        st.error("起始月份不能晚于结束月份"); return

    _mask = (mdf["_ym"].astype(str) >= start_m) & (mdf["_ym"].astype(str) <= end_m)
    fdf      = mdf[_mask].copy()
    orig_fdf = orig_mdf[_mask].copy()
    if len(fdf) == 0:
        st.warning("所选范围内无数据"); return

    # ── 按月聚合 ──
    grp = fdf.groupby("_ym")[metric_cols]
    agg_sum  = grp.sum().reset_index()
    agg_mean = grp.mean().reset_index()
    agg_sum["_ym"]  = agg_sum["_ym"].astype(str)
    agg_mean["_ym"] = agg_mean["_ym"].astype(str)

    # 最新月 vs 上月
    latest_sum  = agg_sum.iloc[-1]
    prev_sum    = agg_sum.iloc[-2]  if len(agg_sum)  >= 2 else None
    latest_mean = agg_mean.iloc[-1]
    prev_mean   = agg_mean.iloc[-2] if len(agg_mean) >= 2 else None

    # ── 趋势图 ──
    st.markdown("---")
    st.markdown("##### 📈 月度趋势")
    tc1, tc2 = st.columns([3, 1])
    with tc1:
        trend_ms = st.multiselect("选择指标", metric_cols,
                                   default=metric_cols[:min(3, len(metric_cols))],
                                   key="m_trend_sel")
    with tc2:
        trend_view = st.radio("维度", ["求和", "均值"], key="m_trend_view")

    tdf = agg_sum if trend_view == "求和" else agg_mean
    if trend_ms:
        _colors = ["#3a7bd5", "#f7971e", "#43a047", "#e53935", "#8e24aa", "#00acc1"]
        fig = _go_m.Figure()
        for idx, m in enumerate(trend_ms):
            fig.add_trace(_go_m.Scatter(
                x=tdf["_ym"], y=tdf[m], name=f"{m}（{trend_view}）",
                mode="lines+markers",
                line=dict(color=_colors[idx % len(_colors)], width=2),
                marker=dict(size=7)
            ))
        fig.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            legend=dict(orientation="h", y=-0.25),
            margin=dict(l=10, r=10, t=10, b=10),
            xaxis=dict(gridcolor="rgba(150,150,150,0.15)", title="月份"),
            yaxis=dict(gridcolor="rgba(150,150,150,0.15)")
        )
        st.plotly_chart(fig, use_container_width=True)

    # ── 月度明细表（每月两行：求和行 + 均值行） ──
    st.markdown("---")
    st.markdown("##### 📋 月度聚合明细")

    def _fmt(v, is_pct=False):
        if not pd.notna(v): return "—"
        suffix = "%" if is_pct else ""
        if v == 0: return f"0{suffix}"
        if abs(v) >= 1 and v == int(v): return f"{int(v):,}{suffix}"
        s = f"{v:,.2f}".rstrip("0").rstrip(".")
        return f"{s}{suffix}"

    rows = []
    for i, ym in enumerate(agg_sum["_ym"]):
        s_row = {"月份": ym, "类型": "求和"}
        a_row = {"月份": ym, "类型": "均值"}
        for m in metric_cols:
            is_p = m in pct_cols
            s_row[m] = _fmt(agg_sum.iloc[i][m], is_p)
            a_row[m] = _fmt(agg_mean.iloc[i][m], is_p)
        rows.append(s_row)
        rows.append(a_row)
    # 合计行 + 区间均值行
    total_row = {"月份": "📌 合计",    "类型": "求和"}
    gavg_row  = {"月份": "📌 区间均值", "类型": "均值"}
    for m in metric_cols:
        is_p = m in pct_cols
        total_row[m] = _fmt(agg_sum[m].sum(), is_p)
        gavg_row[m]  = _fmt(fdf[m].mean(), is_p)
    rows.extend([total_row, gavg_row])
    disp_df = pd.DataFrame(rows)
    st.dataframe(disp_df, use_container_width=True, hide_index=True)

    # ── 原始数据预览 ──
    st.markdown("---")
    st.markdown("##### 🗂 原始数据预览")
    preview_n = st.slider("显示行数", 10, min(200, len(fdf)), min(50, len(fdf)),
                          step=10, key="m_preview_n")
    st.dataframe(
        orig_fdf.head(preview_n).reset_index(drop=True),
        use_container_width=True, hide_index=True
    )

    # ── 下载 ──
    st.markdown("---")
    dl1, dl2 = st.columns(2)
    with dl1:
        exp_df = agg_sum.copy()
        exp_df.columns = ["月份"] + [f"{m}_求和" for m in metric_cols]
        for m in metric_cols:
            exp_df[f"{m}_均值"] = agg_mean[m].values
        _csv_bytes = exp_df.to_csv(index=False, encoding="utf-8-sig").encode("utf-8-sig")
        st.download_button("📥 下载月度聚合 CSV", data=_csv_bytes,
                           file_name=f"月度聚合_{start_m}_{end_m}.csv",
                           mime="text/csv", use_container_width=True, key="m_dl_csv")
    with dl2:
        md_lines = [f"# 月度核心数据报告\n",
                    f"**区间**：{start_m} ～ {end_m}　　**生成时间**：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n",
                    "## 最新月汇总\n"]
        for m in metric_cols:
            sv = latest_sum[m]; mv = latest_mean[m]
            ds = ""
            if prev_sum is not None and pd.notna(prev_sum[m]) and prev_sum[m] != 0:
                dp = (sv - prev_sum[m]) / abs(prev_sum[m]) * 100
                ds = f"（环比 {dp:+.1f}%）"
            md_lines.append(f"- **{m}**：求和 {sv:,.2f} / 均值 {mv:,.2f} {ds}")
        md_text = "\n".join(md_lines)
        st.download_button("📄 下载月报 Markdown", data=md_text.encode("utf-8"),
                           file_name=f"月度报告_{end_m}.md",
                           mime="text/markdown", use_container_width=True, key="m_dl_md")


# ============== 主界面 ==============
def main():
    # ── 认证门控 ──
    if not st.session_state.get("authenticated"):
        show_auth_page()
        return

    uid = st.session_state["user_id"]
    uname = st.session_state["username"]

    st.markdown('''
<div class="main-header">
    <div style="display:flex;align-items:center;justify-content:center;gap:14px">
        <svg width="44" height="44" viewBox="0 0 44 44" fill="none" xmlns="http://www.w3.org/2000/svg">
            <rect width="44" height="44" rx="12" fill="url(#ghdr)"/>
            <rect x="8" y="28" width="6" height="10" rx="2" fill="white" fill-opacity="0.95"/>
            <rect x="17" y="20" width="6" height="18" rx="2" fill="white" fill-opacity="0.95"/>
            <rect x="26" y="14" width="6" height="24" rx="2" fill="white" fill-opacity="0.95"/>
            <polyline points="11,27 20,19 29,13 36,9" stroke="white" stroke-width="2" stroke-opacity="0.5" stroke-linecap="round" fill="none"/>
            <circle cx="36" cy="9" r="2.5" fill="white" fill-opacity="0.9"/>
            <defs><linearGradient id="ghdr" x1="0" y1="0" x2="44" y2="44" gradientUnits="userSpaceOnUse">
                <stop offset="0%" stop-color="#3a7bd5"/><stop offset="100%" stop-color="#00d2ff"/>
            </linearGradient></defs>
        </svg>
        <h1>周均数据对比分析器</h1>
    </div>
    <div class="header-line"></div>
    <p class="header-subtitle">WEEKLY DATA ANALYSIS PLATFORM</p>
</div>
''', unsafe_allow_html=True)
    
    # 侧边栏配置
    with st.sidebar:
        # 用户信息 + 登出
        st.markdown(f"""
        <div style="background:rgba(58,123,213,0.08);border-radius:12px;
                    padding:10px 14px;margin-bottom:12px;
                    border:1px solid rgba(58,123,213,0.18);">
            <span style="font-size:0.8rem;color:#6b7db3;">已登录账户</span><br/>
            <b style="color:#1e3a6e;font-size:1rem;">👤 {uname}</b>
        </div>
        """, unsafe_allow_html=True)
        if st.button("🚪 退出登录", use_container_width=True):
            for k in ["authenticated", "username", "user_id",
                      "comparison_df", "analysis_text", "original_df",
                      "cleaned_df", "process_df", "thresholds",
                      "prev_label", "curr_label"]:
                st.session_state.pop(k, None)
            st.rerun()

        # 深色模式切换
        _is_dark = st.session_state.get("dark_mode", False)
        _theme_label = "☀️ 切换浅色模式" if _is_dark else "🌙 切换深色模式"
        if st.button(_theme_label, use_container_width=True):
            st.session_state["dark_mode"] = not _is_dark
            st.rerun()

        st.divider()
        st.header("⚙️ 配置设置")
        
        # AI设置
        st.subheader("🤖 AI分析设置")
        use_ai = st.checkbox("启用AI智能分析", value=True)
        
        if use_ai:
            api_provider = st.selectbox(
                "API提供商",
                ["OpenRouter", "OpenAI", "自定义"],
                index=0,
                help="选择API提供商"
            )
            
            api_key = st.text_input("API Key", type="password", 
                                    help="输入你的API密钥")
            
            if api_provider == "OpenRouter":
                base_url = "https://openrouter.ai/api/v1"
                selected_model = st.selectbox(
                    "选择模型",
                    [
                        "openai/gpt-4o-mini",
                        "openai/gpt-4o",
                        "anthropic/claude-3.5-sonnet",
                        "anthropic/claude-3-haiku",
                        "google/gemini-2.0-flash-exp:free",
                        "google/gemini-pro",
                        "meta-llama/llama-3.1-70b-instruct",
                        "deepseek/deepseek-chat",
                        "qwen/qwen-2.5-72b-instruct",
                    ],
                    index=0,
                    help="选择OpenRouter上的模型",
                    key="selected_model"
                )
                st.caption("ℹ️ 使用OpenRouter API")
            elif api_provider == "OpenAI":
                base_url = "https://api.openai.com/v1"
                st.caption("ℹ️ 使用OpenAI API")
            else:
                base_url = st.text_input("API Base URL", 
                                        placeholder="https://api.example.com/v1",
                                        help="输入自定义API的Base URL")
        else:
            api_key = ""
            base_url = ""
        
        st.divider()
        
        # 默认阈值
        st.subheader("📏 默认异常阈值")
        use_default_threshold = st.checkbox("启用默认阈值", value=True, 
                                            help="关闭后将使用各指标单独设置的阈值")
        if use_default_threshold:
            default_threshold = st.slider("默认阈值 (%)", 0, 100, 15, 
                                          help="涨跌率超过此值视为异常")
        else:
            default_threshold = 15  # 备用值，但不会被使用
    
    # ── 主功能导航 Tab ──
    _nav = st.radio(
        "", ["📊 周均对比分析", "📅 月度核心看板"],
        horizontal=True, key="_main_nav",
        label_visibility="collapsed"
    )
    if _nav == "📅 月度核心看板":
        show_monthly_dashboard(uid)
        return

    # 主内容区
    st.subheader("📁 数据来源")
    
    # 报表系统快捷入口
    with st.expander("🔗 报表系统快捷入口", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            st.link_button("🚀 打开 Infoc 报表系统", "http://base.cmcm.com:8080/", use_container_width=True)
        with col2:
            custom_url = st.text_input("自定义报表地址", placeholder="输入其他报表系统URL")
            if custom_url:
                st.link_button("打开自定义地址", custom_url, use_container_width=True)
    
    # 快速计算工具
    with st.expander("🧮 快速计算工具（双组对比）", expanded=False):
        st.markdown("**从报表复制数据，粘贴到下方进行快速计算和对比**")
        
        col_a, col_b = st.columns(2)
        
        with col_a:
            st.markdown("##### 🟢 组A")
            input_a = st.text_area("粘贴组A数据", height=100, key="calc_input_a", 
                                   placeholder="从报表复制一列数据粘贴到这里...")
        
        with col_b:
            st.markdown("##### 🟠 组B")
            input_b = st.text_area("粘贴组B数据", height=100, key="calc_input_b",
                                   placeholder="从报表复制另一列数据粘贴到这里...")
        
        def parse_numbers(text):
            """从文本中提取所有数字"""
            import re
            if not text:
                return []
            # 匹配数字（包括带逗号的）
            matches = re.findall(r'-?[\d,]+\.?\d*', text)
            nums = []
            for m in matches:
                try:
                    n = float(m.replace(',', ''))
                    if n != 0:  # 忽略0
                        nums.append(n)
                except:
                    pass
            return nums
        
        def calc_stats(nums):
            """计算统计信息"""
            if not nums:
                return None
            import numpy as np
            return {
                'count': len(nums),
                'sum': sum(nums),
                'avg': sum(nums) / len(nums),
                'max': max(nums),
                'min': min(nums),
                'median': float(np.median(nums))
            }
        
        nums_a = parse_numbers(input_a)
        nums_b = parse_numbers(input_b)
        stats_a = calc_stats(nums_a)
        stats_b = calc_stats(nums_b)
        
        # 显示结果
        result_col_a, result_col_b = st.columns(2)
        
        with result_col_a:
            if stats_a:
                st.success(f"""
                **组A ({stats_a['count']}个数值)**
                - ➕ 求和: **{stats_a['sum']:,.2f}**
                - 📊 均值: **{stats_a['avg']:,.2f}**
                - ⬆️ 最大: **{stats_a['max']:,.2f}**
                - ⬇️ 最小: **{stats_a['min']:,.2f}**
                - 📍 中位数: **{stats_a['median']:,.2f}**
                """)
            else:
                st.info("组A: 未输入数据")
        
        with result_col_b:
            if stats_b:
                st.warning(f"""
                **组B ({stats_b['count']}个数值)**
                - ➕ 求和: **{stats_b['sum']:,.2f}**
                - 📊 均值: **{stats_b['avg']:,.2f}**
                - ⬆️ 最大: **{stats_b['max']:,.2f}**
                - ⬇️ 最小: **{stats_b['min']:,.2f}**
                - 📍 中位数: **{stats_b['median']:,.2f}**
                """)
            else:
                st.info("组B: 未输入数据")
        
        # 对比分析
        if stats_a and stats_b:
            sum_diff = stats_b['sum'] - stats_a['sum']
            sum_pct = (sum_diff / abs(stats_a['sum']) * 100) if stats_a['sum'] != 0 else 0
            avg_diff = stats_b['avg'] - stats_a['avg']
            avg_pct = (avg_diff / abs(stats_a['avg']) * 100) if stats_a['avg'] != 0 else 0
            
            st.markdown("---")
            st.markdown("### 📊 对比分析 (B vs A)")
            
            comp_col1, comp_col2, comp_col3 = st.columns(3)
            with comp_col1:
                st.metric("求和差值", f"{sum_diff:+,.2f}", f"{sum_pct:+.2f}%")
            with comp_col2:
                st.metric("均值差值", f"{avg_diff:+,.2f}", f"{avg_pct:+.2f}%")
            with comp_col3:
                # 涨跌比（B相对于A的变化率）
                change_rate = (stats_b['sum'] / stats_a['sum'] - 1) * 100 if stats_a['sum'] != 0 else 0
                direction = "📈 上涨" if change_rate >= 0 else "📉 下跌"
                st.metric("涨跌比", f"{direction}", f"{change_rate:+.2f}%")
    
    # 截图识别工具
    with st.expander("📷 截图识别数字", expanded=False):
        st.markdown("**上传截图自动识别数字 → 支持双组对比**")
        st.link_button("🚀 打开截图识别工具", "https://shiwassu-tomato-ocr.hf.space", use_container_width=True)
        st.components.v1.iframe("https://shiwassu-tomato-ocr.hf.space", height=620, scrolling=True)
    
    # 数据导入方式选择
    data_source = st.radio(
        "选择数据导入方式",
        ["上传文件", "粘贴数据"],
        horizontal=True,
        help="上传文件支持CSV/Excel；粘贴数据支持从报表系统复制的表格"
    )
    
    original_df = None
    
    if data_source == "上传文件":
        uploaded_file = st.file_uploader(
            "选择数据文件",
            type=['csv', 'xlsx', 'xls'],
            help="支持CSV、Excel格式"
        )
    
        if uploaded_file is not None:
            # 读取文件数据
            try:
                if uploaded_file.name.endswith('.csv'):
                    for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-8-sig']:
                        try:
                            uploaded_file.seek(0)
                            original_df = pd.read_csv(uploaded_file, encoding=encoding)
                            break
                        except:
                            continue
                else:
                    original_df = pd.read_excel(uploaded_file)
                
                st.success(f"✅ 成功加载: {len(original_df)} 行, {len(original_df.columns)} 列")

                # ── 自动匹配阈值模板 ──
                matched = db_auth.find_matching_profile(uid, uploaded_file.name)
                if matched:
                    _key = f"_profile_prompt_{matched['profile_name']}"
                    if not st.session_state.get(_key):
                        st.info(
                            f"📂 检测到文件名「**{uploaded_file.name}**」与已保存模板"
                            f"「**{matched['profile_name']}**」匹配，是否直接应用该阈值配置？"
                        )
                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("✅ 应用模板", key="apply_matched_profile",
                                         use_container_width=True, type="primary"):
                                st.session_state["active_thresholds"] = matched["thresholds"]
                                st.session_state[_key] = True
                                st.success(f"已应用模板「{matched['profile_name']}」")
                                st.rerun()
                        with c2:
                            if st.button("❌ 不应用", key="skip_matched_profile",
                                         use_container_width=True):
                                st.session_state[_key] = True

            except Exception as e:
                st.error(f"文件读取失败: {str(e)}")
                return
    
    else:  # 粘贴数据
        st.markdown("**从报表系统复制表格数据后，粘贴到下方：**")
        pasted_data = st.text_area(
            "粘贴数据",
            height=200,
            placeholder="在报表系统中选择数据区域，Ctrl+C复制，然后在此处Ctrl+V粘贴...\n\n支持Excel/网页表格复制的制表符分隔数据",
            help="支持从Excel、网页表格复制的数据，自动识别制表符或逗号分隔"
        )
        
        if pasted_data:
            try:
                # 尝试解析粘贴的数据
                from io import StringIO
                
                # 先尝试制表符分隔（Excel/网页表格复制）
                if '\t' in pasted_data:
                    original_df = pd.read_csv(StringIO(pasted_data), sep='\t')
                # 再尝试逗号分隔
                elif ',' in pasted_data:
                    original_df = pd.read_csv(StringIO(pasted_data), sep=',')
                else:
                    # 尝试自动检测
                    original_df = pd.read_csv(StringIO(pasted_data), sep=None, engine='python')
                
                st.success(f"✅ 成功解析: {len(original_df)} 行, {len(original_df.columns)} 列")
                
            except Exception as e:
                st.error(f"数据解析失败: {str(e)}\n请确保复制的是完整的表格数据")
                return
    
    if original_df is not None:
        # 侧边栏悬浮计算面板
        with st.sidebar:
            st.markdown("## 🧮 快速计算")
            st.markdown("从下方表格复制数据粘贴到这里")
            
            input_a = st.text_area("🟢 组A数据", height=80, key="sidebar_input_a", 
                                   placeholder="粘贴数据...")
            input_b = st.text_area("🟠 组B数据", height=80, key="sidebar_input_b",
                                   placeholder="粘贴数据...")
            
            def parse_nums(text):
                import re
                if not text: return []
                matches = re.findall(r'-?[\d,]+\.?\d*', text)
                return [float(m.replace(',','')) for m in matches if m and float(m.replace(',','')) != 0]
            
            def show_stats(nums, name):
                if not nums:
                    return None
                import numpy as np
                s = sum(nums)
                return {
                    'count': len(nums), 'sum': s, 'avg': s/len(nums),
                    'max': max(nums), 'min': min(nums), 'median': float(np.median(nums))
                }
            
            nums_a = parse_nums(input_a)
            nums_b = parse_nums(input_b)
            stats_a = show_stats(nums_a, "A")
            stats_b = show_stats(nums_b, "B")
            
            if stats_a:
                st.success(f"""**🟢 组A ({stats_a['count']}个)**
➕ 和: {stats_a['sum']:,.2f}
📊 均: {stats_a['avg']:,.2f}
⬆️ 大: {stats_a['max']:,.0f} ⬇️ 小: {stats_a['min']:,.0f}""")
            
            if stats_b:
                st.warning(f"""**🟠 组B ({stats_b['count']}个)**
➕ 和: {stats_b['sum']:,.2f}
📊 均: {stats_b['avg']:,.2f}
⬆️ 大: {stats_b['max']:,.0f} ⬇️ 小: {stats_b['min']:,.0f}""")
            
            if stats_a and stats_b:
                diff = stats_b['sum'] - stats_a['sum']
                pct = (diff / abs(stats_a['sum']) * 100) if stats_a['sum'] else 0
                st.info(f"""**📊 对比 B-A**
差值: {diff:+,.2f}
变化: {pct:+.2f}%""")
        
        # 原始数据预览
        with st.expander("📋 原始数据预览", expanded=True):
            st.dataframe(original_df, use_container_width=True)
        
        # 数据清洗
        cleaned_df = clean_data(original_df)
        cols = cleaned_df.columns.tolist()
        
        # ========== 数据处理模式选择 ==========
        st.subheader("📊 数据处理模式")
        
        data_mode = st.radio(
            "选择数据处理方式",
            ["模式1: 直接对比（已有周均数据）", "模式2: 日期范围计算（按日期筛选计算均值）", "模式3: 任意两天对比"],
            horizontal=True,
            help="模式1适用于已有周均值数据；模式2适用于按日期范围计算均值；模式3适用于任意选择两天直接对比"
        )
        
        if data_mode == "模式1: 直接对比（已有周均数据）":
            # ========== 模式1: 直接对比 ==========
            st.markdown("##### 🔧 列配置")
            col1, col2, col3 = st.columns(3)
            with col1:
                metric_col = st.selectbox("指标名称列", cols, index=0)
            with col2:
                prev_col = st.selectbox("对比期数据列（如上周）", cols, index=min(1, len(cols)-1))
            with col3:
                curr_col = st.selectbox("当前期数据列（如本周）", cols, index=min(2, len(cols)-1))
            
            prev_label = st.text_input("对比期名称", value="上周均值", key="prev_label_1")
            curr_label = st.text_input("当前期名称", value="本周均值", key="curr_label_1")
            
            process_df = cleaned_df
            date_col = None
            
        elif data_mode == "模式2: 日期范围计算（按日期筛选计算均值）":
            # ========== 模式2: 日期范围计算 ==========
            st.markdown("##### 🔧 列配置")
            
            # 获取可选指标列（排除日期列）
            col1, col2 = st.columns(2)
            with col1:
                date_col = st.selectbox("日期列", cols, index=0, help="选择包含日期的列")
            
            available_metrics = [c for c in cols if c != date_col]
            total_metrics = len(available_metrics)
            
            with col2:
                # 添加标签与日期列对齐
                st.markdown('<p style="font-size: 14px; margin-bottom: 0.5rem;">选择指标</p>', unsafe_allow_html=True)
                
                # 使用唯一的version key来强制刷新checkbox
                if 'checkbox_version' not in st.session_state:
                    st.session_state['checkbox_version'] = 0
                
                # 全选复选框
                if 'select_all_state' not in st.session_state:
                    st.session_state['select_all_state'] = True
                
                def toggle_select_all():
                    new_val = not st.session_state['select_all_state']
                    st.session_state['select_all_state'] = new_val
                    # 增加版本号，强制所有checkbox使用新key重新渲染
                    st.session_state['checkbox_version'] += 1
                    # 更新所有指标选中状态
                    st.session_state['metric_sel_states'] = {m: new_val for m in available_metrics}
                
                # 初始化指标选中状态
                if 'metric_sel_states' not in st.session_state:
                    st.session_state['metric_sel_states'] = {m: True for m in available_metrics}
                
                version = st.session_state['checkbox_version']
                
                # 使用popover实现悬浮下拉框，与日期列平行对齐
                with st.popover("点击选择指标 ▼", use_container_width=True):
                    # 全选复选框放在popover内部顶部
                    st.checkbox(
                        "全选所有指标列", 
                        value=st.session_state['select_all_state'],
                        key="select_all_toggle",
                        on_change=toggle_select_all
                    )
                    st.divider()
                    
                    # 使用容器限制高度并添加滚动条
                    with st.container(height=250):
                        selected = []
                        for metric in available_metrics:
                            # 从session state获取当前状态
                            current_val = st.session_state['metric_sel_states'].get(metric, st.session_state['select_all_state'])
                            # 使用版本号作为key的一部分，确保全选变化时checkbox重新渲染
                            is_checked = st.checkbox(
                                metric, 
                                value=current_val,
                                key=f"m_{version}_{metric}"
                            )
                            # 更新session state
                            st.session_state['metric_sel_states'][metric] = is_checked
                            if is_checked:
                                selected.append(metric)
                        
                        metric_cols_select = selected
                    
                    # 在checkbox渲染后显示实际选中数量
                    st.caption(f"已选择 {len(selected)}/{total_metrics} 个指标")

            # 指标模板管理（Mode 2）
            with st.expander("📋 指标模板管理", expanded=False):
                _mt2_list = db_auth.get_metric_templates(uid)
                _mt2_names = [t["template_name"] for t in _mt2_list]
                if _mt2_list:
                    st.markdown("**已保存的模板**")
                    _mt2_sel = st.selectbox("选择模板", _mt2_names, key="mt2_select")
                    _mt2_data = next((t for t in _mt2_list if t["template_name"] == _mt2_sel), None)
                    if _mt2_data:
                        _m2_prev = _mt2_data["metrics"]
                        st.caption(f"包含 {len(_m2_prev)} 个指标：{', '.join(_m2_prev[:6])}{'…' if len(_m2_prev)>6 else ''}")
                        st.caption(f"📅 更新于：{_mt2_data['updated_at']}")
                    _mt2_c1, _mt2_c2, _mt2_c3 = st.columns(3)
                    with _mt2_c1:
                        if st.button("📥 加载", use_container_width=True, key="mt2_load"):
                            _loaded = set(_mt2_data["metrics"])
                            st.session_state['metric_sel_states'] = {m: (m in _loaded) for m in available_metrics}
                            st.session_state['select_all_state'] = False
                            st.session_state['checkbox_version'] += 1
                            st.success(f"已加载「{_mt2_sel}」")
                            st.rerun()
                    with _mt2_c2:
                        _mt2_rn = st.text_input("重命名为", key="mt2_rename_input", placeholder="输入新名称")
                        if st.button("✏️ 重命名", use_container_width=True, key="mt2_rename"):
                            ok, msg = db_auth.rename_metric_template(uid, _mt2_sel, _mt2_rn)
                            st.success(msg) if ok else st.error(msg)
                            st.rerun()
                    with _mt2_c3:
                        if st.button("🗑️ 删除", use_container_width=True, key="mt2_del"):
                            ok, msg = db_auth.delete_metric_template(uid, _mt2_sel)
                            st.success(msg) if ok else st.error(msg)
                            st.rerun()
                else:
                    st.caption("暂无已保存的指标模板")
                st.divider()
                st.markdown("**保存当前选中指标为模板**")
                _mt2_save_name = st.text_input("模板名称", key="mt2_save_name", placeholder="例：核心业务指标")
                if st.button("💾 保存当前指标", use_container_width=True, type="primary", key="mt2_save"):
                    _cur2 = [m for m in available_metrics if st.session_state['metric_sel_states'].get(m, True)]
                    ok, msg = db_auth.save_metric_template(uid, _mt2_save_name, _cur2)
                    st.success(msg) if ok else st.error(msg)
                    if ok:
                        st.rerun()

            if not metric_cols_select:
                st.warning("请至少选择一个指标列")
                return

            # 解析日期获取范围（支持带星期标注的日期如 2026-01-11(日)）
            try:
                cleaned_df[date_col] = cleaned_df[date_col].apply(parse_date_with_weekday)
                min_date = cleaned_df[date_col].min().date()
                max_date = cleaned_df[date_col].max().date()
            except:
                st.error("日期列解析失败，请检查日期格式")
                return

            st.markdown("##### 📅 日期范围选择")

            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**对比期（如上周）**")
                prev_start = st.date_input("开始日期", value=min_date, key="prev_start")
                prev_end = st.date_input("结束日期", value=min_date + timedelta(days=6), key="prev_end")
                prev_label = st.text_input("对比期名称", value="上周均值", key="prev_label_2")

            with col2:
                st.markdown("**当前期（如本周）**")
                curr_start = st.date_input("开始日期", value=max_date - timedelta(days=6), key="curr_start")
                curr_end = st.date_input("结束日期", value=max_date, key="curr_end")
                curr_label = st.text_input("当前期名称", value="本周均值", key="curr_label_2")

            # 计算两个周期的均值（传入原始数据以检测格式）
            prev_avg, prev_records, format_info, prev_range_days = calculate_daily_average(
                cleaned_df, date_col, metric_cols_select, prev_start, prev_end, original_df
            )
            curr_avg, curr_records, _, curr_range_days = calculate_daily_average(
                cleaned_df, date_col, metric_cols_select, curr_start, curr_end, original_df
            )

            if prev_avg is None or curr_avg is None:
                st.error("所选日期范围内没有数据，请调整日期范围")
                return

            st.info(f"📊 对比期: {prev_start} ~ {prev_end} ({prev_range_days}天, {prev_records}条记录) | 当前期: {curr_start} ~ {curr_end} ({curr_range_days}天, {curr_records}条记录)")

            # 构建对比数据（保持原始格式）
            process_data = []
            for col_name in metric_cols_select:
                prev_val = prev_avg.get(col_name, 0)
                curr_val = curr_avg.get(col_name, 0)
                col_format = format_info.get(col_name, {'is_percent': False, 'decimal_places': 0})
                process_data.append({
                    '指标': col_name,
                    prev_label: prev_val,
                    curr_label: curr_val,
                    '_format': col_format
                })
            process_df = pd.DataFrame(process_data)

            st.session_state['format_info'] = format_info

            metric_col = '指标'
            prev_col = prev_label
            curr_col = curr_label

            # 显示趋势图
            st.markdown("##### 📈 数据趋势图")
            trend_metrics = st.multiselect(
                "选择要显示趋势的指标",
                metric_cols_select,
                default=metric_cols_select[:3] if len(metric_cols_select) >= 3 else metric_cols_select,
                key="trend_metrics"
            )

            if trend_metrics:
                trend_fig = create_trend_chart(cleaned_df, date_col, trend_metrics, "每日数据趋势")
                st.plotly_chart(trend_fig, use_container_width=True)
        
        elif data_mode == "模式3: 任意两天对比":
            # ========== 模式3: 任意两天对比 ==========
            st.markdown("##### 🔧 列配置")
            
            col1, col2 = st.columns(2)
            with col1:
                date_col = st.selectbox("日期列", cols, index=0, help="选择包含日期的列", key="mode3_date_col")
            
            available_metrics = [c for c in cols if c != date_col]
            total_metrics = len(available_metrics)
            
            with col2:
                st.markdown('<p style="font-size: 14px; margin-bottom: 0.5rem;">选择指标</p>', unsafe_allow_html=True)
                
                if 'checkbox_version_m3' not in st.session_state:
                    st.session_state['checkbox_version_m3'] = 0
                
                if 'select_all_state_m3' not in st.session_state:
                    st.session_state['select_all_state_m3'] = True
                
                def toggle_select_all_m3():
                    new_val = not st.session_state['select_all_state_m3']
                    st.session_state['select_all_state_m3'] = new_val
                    st.session_state['checkbox_version_m3'] += 1
                    st.session_state['metric_sel_states_m3'] = {m: new_val for m in available_metrics}
                
                if 'metric_sel_states_m3' not in st.session_state:
                    st.session_state['metric_sel_states_m3'] = {m: True for m in available_metrics}
                
                version_m3 = st.session_state['checkbox_version_m3']
                
                with st.popover("点击选择指标 ▼", use_container_width=True):
                    st.checkbox(
                        "全选所有指标列", 
                        value=st.session_state['select_all_state_m3'],
                        key="select_all_toggle_m3",
                        on_change=toggle_select_all_m3
                    )
                    st.divider()
                    
                    with st.container(height=250):
                        selected_m3 = []
                        for metric in available_metrics:
                            current_val = st.session_state['metric_sel_states_m3'].get(metric, st.session_state['select_all_state_m3'])
                            is_checked = st.checkbox(
                                metric, 
                                value=current_val,
                                key=f"m3_{version_m3}_{metric}"
                            )
                            st.session_state['metric_sel_states_m3'][metric] = is_checked
                            if is_checked:
                                selected_m3.append(metric)
                        
                        metric_cols_select = selected_m3
                    
                    st.caption(f"已选择 {len(selected_m3)}/{total_metrics} 个指标")

            # 指标模板管理（Mode 3）
            with st.expander("📋 指标模板管理", expanded=False):
                _mt3_list = db_auth.get_metric_templates(uid)
                _mt3_names = [t["template_name"] for t in _mt3_list]
                if _mt3_list:
                    st.markdown("**已保存的模板**")
                    _mt3_sel = st.selectbox("选择模板", _mt3_names, key="mt3_select")
                    _mt3_data = next((t for t in _mt3_list if t["template_name"] == _mt3_sel), None)
                    if _mt3_data:
                        _m3_prev = _mt3_data["metrics"]
                        st.caption(f"包含 {len(_m3_prev)} 个指标：{', '.join(_m3_prev[:6])}{'…' if len(_m3_prev)>6 else ''}")
                        st.caption(f"📅 更新于：{_mt3_data['updated_at']}")
                    _mt3_c1, _mt3_c2, _mt3_c3 = st.columns(3)
                    with _mt3_c1:
                        if st.button("📥 加载", use_container_width=True, key="mt3_load"):
                            _loaded3 = set(_mt3_data["metrics"])
                            st.session_state['metric_sel_states_m3'] = {m: (m in _loaded3) for m in available_metrics}
                            st.session_state['select_all_state_m3'] = False
                            st.session_state['checkbox_version_m3'] += 1
                            st.success(f"已加载「{_mt3_sel}」")
                            st.rerun()
                    with _mt3_c2:
                        _mt3_rn = st.text_input("重命名为", key="mt3_rename_input", placeholder="输入新名称")
                        if st.button("✏️ 重命名", use_container_width=True, key="mt3_rename"):
                            ok, msg = db_auth.rename_metric_template(uid, _mt3_sel, _mt3_rn)
                            st.success(msg) if ok else st.error(msg)
                            st.rerun()
                    with _mt3_c3:
                        if st.button("🗑️ 删除", use_container_width=True, key="mt3_del"):
                            ok, msg = db_auth.delete_metric_template(uid, _mt3_sel)
                            st.success(msg) if ok else st.error(msg)
                            st.rerun()
                else:
                    st.caption("暂无已保存的指标模板")
                st.divider()
                st.markdown("**保存当前选中指标为模板**")
                _mt3_save_name = st.text_input("模板名称", key="mt3_save_name", placeholder="例：核心业务指标")
                if st.button("💾 保存当前指标", use_container_width=True, type="primary", key="mt3_save"):
                    _cur3 = [m for m in available_metrics if st.session_state.get('metric_sel_states_m3', {}).get(m, True)]
                    ok, msg = db_auth.save_metric_template(uid, _mt3_save_name, _cur3)
                    st.success(msg) if ok else st.error(msg)
                    if ok:
                        st.rerun()

            if not metric_cols_select:
                st.warning("请至少选择一个指标列")
                return

            # 解析日期
            try:
                cleaned_df[date_col] = cleaned_df[date_col].apply(parse_date_with_weekday)
                available_dates = sorted(cleaned_df[date_col].dropna().unique())
                available_dates_str = [d.strftime('%Y-%m-%d') if hasattr(d, 'strftime') else str(d) for d in available_dates]
            except:
                st.error("日期列解析失败，请检查日期格式")
                return

            st.markdown("##### 📅 选择对比日期")

            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**日期1（对比基准）**")
                if len(available_dates_str) >= 2:
                    day1_idx = st.selectbox("选择日期1", range(len(available_dates_str)),
                                           format_func=lambda x: available_dates_str[x],
                                           index=len(available_dates_str)-2,
                                           key="day1_select")
                else:
                    day1_idx = st.selectbox("选择日期1", range(len(available_dates_str)),
                                           format_func=lambda x: available_dates_str[x],
                                           index=0,
                                           key="day1_select")
                prev_label = st.text_input("日期1名称", value=available_dates_str[day1_idx], key="day1_label")

            with col2:
                st.markdown("**日期2（当前对比）**")
                day2_idx = st.selectbox("选择日期2", range(len(available_dates_str)),
                                       format_func=lambda x: available_dates_str[x],
                                       index=len(available_dates_str)-1,
                                       key="day2_select")
                curr_label = st.text_input("日期2名称", value=available_dates_str[day2_idx], key="day2_label")

            # 获取两天的数据
            day1_date = available_dates[day1_idx]
            day2_date = available_dates[day2_idx]

            day1_data = cleaned_df[cleaned_df[date_col] == day1_date]
            day2_data = cleaned_df[cleaned_df[date_col] == day2_date]

            if day1_data.empty or day2_data.empty:
                st.error("所选日期没有数据，请重新选择")
                return

            st.info(f"📊 对比: {available_dates_str[day1_idx]} vs {available_dates_str[day2_idx]}")

            # 检测原始数据格式
            format_info = {}
            for col_name in metric_cols_select:
                if col_name in original_df.columns:
                    sample_values = original_df[col_name].dropna().head(10)
                    is_percent = False
                    decimal_places = 0
                    for val in sample_values:
                        if isinstance(val, str) and '%' in val:
                            is_percent = True
                            try:
                                num_str = val.replace('%', '').replace(',', '').strip()
                                if '.' in num_str:
                                    decimal_places = max(decimal_places, len(num_str.split('.')[-1]))
                            except:
                                pass
                            break
                    format_info[col_name] = {'is_percent': is_percent, 'decimal_places': decimal_places}
                else:
                    format_info[col_name] = {'is_percent': False, 'decimal_places': 0}

            # 构建对比数据
            process_data = []
            for col_name in metric_cols_select:
                prev_val = day1_data[col_name].iloc[0] if col_name in day1_data.columns else 0
                curr_val = day2_data[col_name].iloc[0] if col_name in day2_data.columns else 0

                col_format = format_info.get(col_name, {'is_percent': False, 'decimal_places': 0})
                # 与Mode2保持一致：百分比值乘以100，使差值和格式化输出正确
                prev_val = format_value(prev_val, col_format)
                curr_val = format_value(curr_val, col_format)

                process_data.append({
                    '指标': col_name,
                    prev_label: prev_val,
                    curr_label: curr_val,
                    '_format': col_format
                })
            process_df = pd.DataFrame(process_data)

            st.session_state['format_info'] = format_info

            metric_col = '指标'
            prev_col = prev_label
            curr_col = curr_label

            # 显示趋势图
            st.markdown("##### 📈 数据趋势图")
            trend_metrics = st.multiselect(
                "选择要显示趋势的指标",
                metric_cols_select,
                default=metric_cols_select[:3] if len(metric_cols_select) >= 3 else metric_cols_select,
                key="trend_metrics_m3"
            )

            if trend_metrics:
                trend_fig = create_trend_chart(cleaned_df, date_col, trend_metrics, "每日数据趋势")
                st.plotly_chart(trend_fig, use_container_width=True)
        
        # 自定义阈值设置
        st.subheader("⚡ 异常阈值设定")
        st.caption("为各指标设置异常阈值（涨跌率超过阈值视为异常）")
        
        # 初始化阈值（支持从已加载模板读取）
        _active = st.session_state.get("active_thresholds", {})
        thresholds = {'__default__': _active.get('__default__', default_threshold)}
        
        # 获取所有指标（根据模式不同获取方式不同）
        if data_mode == "模式1: 直接对比（已有周均数据）":
            import re as _re_m1
            _dpat = _re_m1.compile(r'^\d{4}[-/]\d{1,2}[-/]\d{1,2}')
            def _vals_are_date_like(col):
                sample = cleaned_df[col].dropna().astype(str).head(10)
                hits = sum(1 for v in sample if _dpat.match(v.strip()))
                return hits >= max(1, len(sample) // 2)
            def _s_is_date_or_num(s):
                if _dpat.match(s): return True
                try: float(s.replace(',', '')); return True
                except ValueError: return False
            _raw = process_df[metric_col].dropna().unique().tolist()
            _raw_strs = [str(v).strip() for v in _raw[:10]]
            if _raw_strs and all(_s_is_date_or_num(s) for s in _raw_strs):
                # metric_col 无真正指标名 → 用列名作为指标，按值排除日期列
                all_metrics = [c for c in cols if not _vals_are_date_like(c)]
            else:
                all_metrics = _raw
        else:
            # 模式2/3: 从选择的指标列名获取
            all_metrics = metric_cols_select
        
        # 选择要分析的指标（带全选功能）
        st.markdown("##### 📋 选择分析指标")
        if st.session_state.pop("_mt1_uncheck", False):
            st.session_state["select_all_metrics"] = False
        select_all = st.checkbox("全选所有指标", value=True, key="select_all_metrics")

        if select_all:
            selected_metrics = all_metrics
        else:
            _m1_default = st.session_state.pop("_m1_tpl_load", all_metrics)
            selected_metrics = st.multiselect(
                "选择要分析的指标",
                all_metrics,
                default=[m for m in _m1_default if m in all_metrics] or all_metrics,
                key="metric_multiselect",
                help="选择需要进行分析的指标"
            )

        # 指标模板管理（Mode 1）
        with st.expander("📋 指标模板管理", expanded=False):
            _mt1_list = db_auth.get_metric_templates(uid)
            _mt1_names = [t["template_name"] for t in _mt1_list]
            if _mt1_list:
                st.markdown("**已保存的模板**")
                _mt1_sel = st.selectbox("选择模板", _mt1_names, key="mt1_select")
                _mt1_data = next((t for t in _mt1_list if t["template_name"] == _mt1_sel), None)
                if _mt1_data:
                    _m1_prev = _mt1_data["metrics"]
                    st.caption(f"包含 {len(_m1_prev)} 个指标：{', '.join(_m1_prev[:6])}{'…' if len(_m1_prev)>6 else ''}")
                    st.caption(f"📅 更新于：{_mt1_data['updated_at']}")
                _mt1_c1, _mt1_c2, _mt1_c3 = st.columns(3)
                with _mt1_c1:
                    if st.button("📥 加载", use_container_width=True, key="mt1_load"):
                        st.session_state["_mt1_uncheck"] = True
                        st.session_state["_m1_tpl_load"] = _mt1_data["metrics"]
                        st.success(f"已加载「{_mt1_sel}」")
                        st.rerun()
                with _mt1_c2:
                    _mt1_rn = st.text_input("重命名为", key="mt1_rename_input", placeholder="输入新名称")
                    if st.button("✏️ 重命名", use_container_width=True, key="mt1_rename"):
                        ok, msg = db_auth.rename_metric_template(uid, _mt1_sel, _mt1_rn)
                        st.success(msg) if ok else st.error(msg)
                        st.rerun()
                with _mt1_c3:
                    if st.button("🗑️ 删除", use_container_width=True, key="mt1_del"):
                        ok, msg = db_auth.delete_metric_template(uid, _mt1_sel)
                        st.success(msg) if ok else st.error(msg)
                        st.rerun()
            else:
                st.caption("暂无已保存的指标模板")
            st.divider()
            st.markdown("**保存当前选中指标为模板**")
            _mt1_save_name = st.text_input("模板名称", key="mt1_save_name", placeholder="例：核心业务指标")
            if st.button("💾 保存当前指标", use_container_width=True, type="primary", key="mt1_save"):
                ok, msg = db_auth.save_metric_template(uid, _mt1_save_name, list(selected_metrics))
                st.success(msg) if ok else st.error(msg)
                if ok:
                    st.rerun()

        if not selected_metrics:
            st.warning("请至少选择一个指标进行分析")
            return

        # 更新 metrics 为选中的指标
        metrics = selected_metrics
        
        # 阈值配置（左边指标名，右边滑动条）
        with st.expander("📏 各指标阈值配置", expanded=True):
            if use_default_threshold:
                st.caption(f"🔒 已启用默认阈值，所有指标使用 {default_threshold}%")
            else:
                st.caption("🔓 可为每个指标单独设置阈值")
            
            for m in metrics:
                col1, col2 = st.columns([1, 3])
                with col1:
                    st.markdown(f"**{m}**")
                with col2:
                    if use_default_threshold:
                        # 启用默认阈值时，锁定为默认值，滑动条禁用
                        thresholds[m] = default_threshold
                        st.slider(
                            f"阈值",
                            min_value=0,
                            max_value=100,
                            value=default_threshold,
                            key=f"slider_{m}",
                            label_visibility="collapsed",
                            disabled=True
                        )
                    else:
                        # 取消默认阈值时，可自由编辑；若有加载的模板则用模板值
                        _init_val = _active.get(m, 15)
                        thresholds[m] = st.slider(
                            f"阈值",
                            min_value=0,
                            max_value=100,
                            value=_init_val,
                            key=f"slider_{m}",
                            label_visibility="collapsed",
                            disabled=False
                        )
        
        # ── 阈值模板管理 ──────────────────────────────────────────
        with st.expander("💾 阈值模板管理", expanded=False):
            _profiles = db_auth.get_profiles(uid)
            _profile_names = [p["profile_name"] for p in _profiles]

            if _profiles:
                st.markdown("**已保存的模板**")
                _sel = st.selectbox("选择模板", _profile_names, key="profile_select")
                _sel_data = next((p for p in _profiles if p["profile_name"] == _sel), None)

                _c1, _c2, _c3 = st.columns(3)
                with _c1:
                    if st.button("📥 加载", use_container_width=True, key="load_profile"):
                        # 清除已有滑块 session_state，使 value= 生效
                        for _k in list(st.session_state.keys()):
                            if _k.startswith("slider_"):
                                del st.session_state[_k]
                        st.session_state["active_thresholds"] = _sel_data["thresholds"]
                        st.success(f"已加载模板「{_sel}」，请查看阈值已更新")
                        st.rerun()
                with _c2:
                    _new_name = st.text_input("重命名为", key="rename_input",
                                             placeholder="输入新名称")
                    if st.button("✏️ 重命名", use_container_width=True, key="rename_profile"):
                        ok, msg = db_auth.rename_profile(uid, _sel, _new_name)
                        st.success(msg) if ok else st.error(msg)
                        st.rerun()
                with _c3:
                    if st.button("🗑️ 删除", use_container_width=True, key="del_profile"):
                        ok, msg = db_auth.delete_profile(uid, _sel)
                        st.success(msg) if ok else st.error(msg)
                        if ok and st.session_state.get("active_thresholds") == _sel_data["thresholds"]:
                            st.session_state.pop("active_thresholds", None)
                        st.rerun()

                if _sel_data:
                    st.caption(f"📅 上次更新：{_sel_data['updated_at']}")
            else:
                st.caption("暂无已保存模板")

            st.divider()
            st.markdown("**保存当前阈值为新模板**")
            _save_name = st.text_input("模板名称", key="save_profile_name",
                                       placeholder="例：会员业务_周报")
            if st.button("💾 保存当前阈值", use_container_width=True,
                         type="primary", key="save_profile_btn"):
                ok, msg = db_auth.save_profile(uid, _save_name, thresholds)
                st.success(msg) if ok else st.error(msg)
                if ok:
                    st.rerun()

        # 执行分析
        if st.button("🚀 开始分析", type="primary", use_container_width=True):
            with st.spinner("正在分析数据..."):
                try:
                    # 获取格式信息
                    current_format_info = st.session_state.get('format_info', {})
                    
                    # 筛选选中的指标
                    # Mode1列名作为指标时，isin匹配不到行，直接用全量数据
                    _cand = process_df[process_df[metric_col].isin(metrics)]
                    filtered_df = (_cand if not _cand.empty else process_df).copy()

                    # Mode1：按行检测百分比格式并将小数值还原为百分比（0.0523→5.23）
                    if data_mode == "模式1: 直接对比（已有周均数据）":
                        m1_fmt = {}
                        if metric_col in original_df.columns:
                            for _, orig_row in original_df.iterrows():
                                m_name = str(orig_row.get(metric_col, '')).strip()
                                if not m_name:
                                    continue
                                fmt = {'is_percent': False, 'decimal_places': 0}
                                for chk_col in [prev_col, curr_col]:
                                    raw_v = str(orig_row.get(chk_col, '')).strip()
                                    if '%' in raw_v:
                                        fmt['is_percent'] = True
                                        num_p = raw_v.replace('%', '').replace(',', '')
                                        if '.' in num_p:
                                            fmt['decimal_places'] = len(num_p.split('.')[-1])
                                        break
                                m1_fmt[m_name] = fmt
                                if fmt['is_percent']:
                                    _mask = filtered_df[metric_col].astype(str).str.strip() == m_name
                                    for _c in [prev_col, curr_col]:
                                        if _c in filtered_df.columns:
                                            filtered_df.loc[_mask, _c] = (
                                                pd.to_numeric(filtered_df.loc[_mask, _c], errors='coerce') * 100
                                            )
                        if m1_fmt:
                            current_format_info = m1_fmt
                    
                    # 计算对比
                    comparison_df = calculate_comparison(
                        filtered_df, metric_col, prev_col, curr_col, thresholds,
                        prev_label=prev_label, curr_label=curr_label,
                        format_info=current_format_info
                    )
                    
                    # 存储到session
                    st.session_state['comparison_df'] = comparison_df
                    st.session_state['original_df'] = original_df
                    st.session_state['cleaned_df'] = cleaned_df
                    st.session_state['process_df'] = process_df
                    st.session_state['thresholds'] = thresholds
                    st.session_state['prev_label'] = prev_label
                    st.session_state['curr_label'] = curr_label
                except Exception as _e:
                    st.error(f"分析出错：{_e}")
                    st.stop()
                
            st.success("✅ 数据分析完成!")
        
    # 示例阈值配置（无真实数据时展示）
    if 'comparison_df' not in st.session_state:
        st.divider()
        st.subheader("⚡ 异常阈值设定（示例）")
        st.caption("上传数据后，可为每个指标单独配置异常阈值。以下为示例效果：")
        with st.expander("📏 各指标阈值配置（示例）", expanded=True):
            st.caption("🔒 示例：已启用默认阈值，所有指标统一使用 15%")
            demo_metrics = ['DAU', 'MAU', '次日留存', '付费率', 'ARPU']
            for m in demo_metrics:
                c1, c2 = st.columns([1, 3])
                with c1:
                    st.markdown(f"**{m}**")
                with c2:
                    st.slider(f"阈值_{m}", 0, 100, 15,
                              label_visibility="collapsed", disabled=True,
                              key=f"demo_slider_{m}")

    # 显示分析结果（无真实数据时显示示例）
    is_demo = 'comparison_df' not in st.session_state
    if is_demo:
        comparison_df = get_demo_comparison()
        prev_lbl, curr_lbl = '上周均值', '本周均值'
        demo_thresholds = {'__default__': 15}
    else:
        comparison_df = st.session_state['comparison_df']
        prev_lbl = st.session_state.get('prev_label', '上周均值')
        curr_lbl = st.session_state.get('curr_label', '本周均值')
        demo_thresholds = st.session_state.get('thresholds', {'__default__': 15})

    st.divider()
    if is_demo:
        st.subheader("📊 分析结果（示例）")
        st.info("📌 以下为示例数据效果预览。上传真实数据并点击「🚀 开始分析」后，将自动更新为实际结果。")
    else:
        st.subheader("📊 分析结果")

    # 统计卡片
    def parse_rate_value(rate_str):
        try:
            return float(str(rate_str).replace('%', '').replace('+', ''))
        except:
            return 0

    rate_numeric = comparison_df['涨跌率(%)'].apply(parse_rate_value)

    total = len(comparison_df)
    abnormal_count = int(comparison_df['是否异常'].sum())
    normal_count = total - abnormal_count
    abnormal_pct = f"{abnormal_count/total*100:.0f}%" if total > 0 else "0%"

    # 解析差值列为数值，用于求和/均值
    def parse_diff(v):
        try:
            return float(str(v).replace('%', '').replace('+', '').replace(',', ''))
        except:
            return 0.0
    diff_numeric = comparison_df['差值'].apply(parse_diff)
    diff_sum = diff_numeric.sum()
    diff_mean = diff_numeric.mean()
    avg_rate = rate_numeric.mean()

    # 行1：指标概览
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("总指标数", total)
    with col2:
        st.metric("正常指标", normal_count)
    with col3:
        st.metric("异常指标", abnormal_count)
    with col4:
        st.metric("异常比例", abnormal_pct)

    # 行2：差值汇总
    col5, col6, col7 = st.columns(3)
    with col5:
        st.metric("求和差值", f"{diff_sum:+,.1f}", delta=f"{avg_rate:+.1f}%")
    with col6:
        st.metric("均值差值", f"{diff_mean:+,.1f}", delta=f"{avg_rate:+.1f}%")
    with col7:
        direction = "📈 上涨" if avg_rate >= 0 else "📉 下跌"
        st.metric("涨跌比", direction, delta=f"{avg_rate:+.2f}%")

    # 对比数据表
    st.markdown("#### 📈 对比数据")

    def highlight_abnormal(row):
        if row['是否异常']:
            return ['background-color: #5c1a1a; color: #ffb3b3'] * len(row)
        return [''] * len(row)

    if st.session_state.get("dark_mode", False):
        def _df_to_dark_html(df):
            th_style = "background:#1c2333;color:#79a8ff;padding:8px 12px;text-align:left;border-bottom:2px solid #30363d;font-weight:600;font-size:0.88rem;"
            td_style = "padding:7px 12px;border-bottom:1px solid #21262d;color:#c9d1d9;font-size:0.88rem;"
            td_abn   = "padding:7px 12px;border-bottom:1px solid #21262d;color:#ffb3b3;font-size:0.88rem;"
            row_norm = "background:#0d1117;"
            row_abn  = "background:#5c1a1a;"
            html = ['<div style="overflow-x:auto;border-radius:10px;border:1px solid #30363d;">',
                    '<table style="width:100%;border-collapse:collapse;">',
                    '<thead><tr>',
                    f'<th style="{th_style}">#</th>']
            for col in df.columns:
                html.append(f'<th style="{th_style}">{col}</th>')
            html.append('</tr></thead><tbody>')
            for i, (_, row) in enumerate(df.iterrows()):
                is_abn = bool(row.get('是否异常', False))
                rs = row_abn if is_abn else row_norm
                td = td_abn  if is_abn else td_style
                html.append(f'<tr style="{rs}">')
                html.append(f'<td style="{td}">{i}</td>')
                for col in df.columns:
                    val = row[col]
                    if col == '是否异常':
                        val = '✅' if val else '☐'
                    html.append(f'<td style="{td}">{val}</td>')
                html.append('</tr>')
            html.append('</tbody></table></div>')
            return ''.join(html)
        st.markdown(_df_to_dark_html(comparison_df), unsafe_allow_html=True)
    else:
        styled_df = comparison_df.style.apply(highlight_abnormal, axis=1)
        st.dataframe(styled_df, use_container_width=True)

    # 对比数据下载
    _dl1, _dl2 = st.columns(2)
    with _dl1:
        _csv = comparison_df.to_csv(index=False, encoding='utf-8-sig')
        st.download_button(
            "📥 下载对比数据 CSV",
            data=_csv.encode('utf-8-sig'),
            file_name=f"对比数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
            use_container_width=True,
            key="dl_cmp_csv"
        )
    with _dl2:
        _xls_buf = BytesIO()
        comparison_df.to_excel(_xls_buf, index=False)
        _xls_buf.seek(0)
        st.download_button(
            "📊 下载对比数据 Excel",
            data=_xls_buf,
            file_name=f"对比数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="dl_cmp_xlsx"
        )

    abnormal_df = comparison_df[comparison_df['是否异常'] == True]
    if not abnormal_df.empty:
        st.markdown("#### ⚠️ 异常数据")
        if st.session_state.get("dark_mode", False):
            st.markdown(_df_to_dark_html(abnormal_df), unsafe_allow_html=True)
        else:
            st.dataframe(abnormal_df, use_container_width=True)

    # 可视化图表
    st.markdown("#### 📊 数据可视化")
    chart_tab1, chart_tab2 = st.tabs(["对比柱状图", "涨跌率分布"])

    with chart_tab1:
        bar_fig = create_comparison_bar_chart(comparison_df, '指标', prev_lbl, curr_lbl)
        st.plotly_chart(bar_fig, use_container_width=True)

    with chart_tab2:
        rate_fig = create_change_rate_chart(comparison_df, demo_thresholds)
        st.plotly_chart(rate_fig, use_container_width=True)

    st.divider()

    # AI分析
    st.subheader("🤖 智能分析报告")
    if is_demo:
        st.caption("上传真实数据并完成分析后，可在此生成AI报告")
        st.button("📝 生成分析报告", use_container_width=True, disabled=True)
    else:
        if st.button("📝 生成分析报告", use_container_width=True):
            with st.spinner("AI正在分析中..."):
                analysis_text = generate_ai_analysis(
                    comparison_df,
                    st.session_state['original_df'],
                    st.session_state['thresholds'],
                    api_key if use_ai else "",
                    base_url if use_ai and base_url else None
                )
                st.session_state['analysis_text'] = analysis_text

        if 'analysis_text' in st.session_state:
            st.markdown(st.session_state['analysis_text'])

            st.divider()

            st.subheader("📥 导出报告")
            col1, col2, col3 = st.columns(3)

            with col1:
                excel_report = create_excel_report(
                    comparison_df,
                    st.session_state['analysis_text'],
                    st.session_state['thresholds']
                )
                st.download_button(
                    label="📊 下载Excel报告",
                    data=excel_report,
                    file_name=f"周均分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

            with col2:
                word_report = create_word_report(
                    comparison_df,
                    st.session_state['analysis_text'],
                    st.session_state['thresholds']
                )
                st.download_button(
                    label="📝 下载Word报告",
                    data=word_report,
                    file_name=f"周均分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            with col3:
                md_report = create_markdown_report(
                    comparison_df,
                    st.session_state['analysis_text'],
                    st.session_state['thresholds']
                )
                st.download_button(
                    label="📄 下载Markdown报告",
                    data=md_report.encode('utf-8'),
                    file_name=f"周均分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown",
                    use_container_width=True
                )

    if is_demo:
        st.divider()
        st.subheader("📥 导出报告")
        st.caption("上传真实数据并生成分析报告后，可在此下载 Excel / Word / Markdown 报告")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.button("📊 下载Excel报告", use_container_width=True, disabled=True)
        with col2:
            st.button("📝 下载Word报告", use_container_width=True, disabled=True)
        with col3:
            st.button("📄 下载Markdown报告", use_container_width=True, disabled=True)


if __name__ == "__main__":
    main()
